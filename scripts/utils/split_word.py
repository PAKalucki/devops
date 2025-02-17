from docx import Document

# Load the existing document
original_doc = Document("DnD55e_2024.docx")

# Create a new document
new_doc = Document()

# Loop through the content and add it to the new document
for i, paragraph in enumerate(original_doc.paragraphs):
    if i + 1 >= 1 and i + 1 <= 41:  # Pages are based on content
        new_doc.add_paragraph(paragraph.text)

# Save the new document
new_doc.save("DnD55e_2024_rules.docx")
